# coding=utf-8
"""
@Time : 2024/1/3
@GROUP : 安全技术部
@AUTH : 李可佳
@EMAIL : security@uniontech.com
@COPY_RIGHT : Copyright (C) 2024 Uniontech Security All Rights Reserved.
"""

import requests
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
from io import BytesIO
from matplotlib.patches import Patch
import global_variable

plt.rcParams['font.sans-serif'] = 'Noto Sans CJK JP'

plt.rcParams.update({'font.size': 12})


class VulInfo:
    def __init__(self, cveid, source, score, level, description, is_fixed, fix_suggest, reference_link, affectPkg=None):
        self.cveid = cveid
        self.source = source
        self.score = score
        self.level = level
        self.description = description
        self.is_fixed = is_fixed
        self.affectPkg = affectPkg
        self.fix_suggest = fix_suggest
        self.reference_link = reference_link


def color_for_cell(cell):
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="B0C4DE"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm_1)


def insert_table(doc, vul_info, index, server_type):
    heading = doc.add_heading("\n" + vul_info.cveid, level=4)
    # 获取标题的 run 对象
    run = heading.runs[0]  # 假设这是第一个 run 对象
    if float(vul_info.score) >= 9:
        # 设置颜色（RGBColor(r, g, b) 中 r、g、b 分别表示红、绿、蓝的值）
        run.font.color.rgb = RGBColor(192, 0, 0)
    elif float(vul_info.score) >= 7:
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif float(vul_info.score) >= 4:
        run.font.color.rgb = RGBColor(247, 150, 70)
    else:
        run.font.color.rgb = RGBColor(255, 192, 0)

    if vul_info.affectPkg is None:
        table = doc.add_table(10, 3, style="Table Grid")
        table.autofit = False
        table.allow_autofit = False
        # 设置每列的宽度，单位为英寸
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(2)
        table.columns[2].width = Inches(2)

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 0).merge(table.cell(2, 2))
        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(3, 0).merge(table.cell(3, 2))

        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(4, 0).merge(table.cell(4, 2))
        table.cell(5, 0).merge(table.cell(5, 1))
        table.cell(5, 0).merge(table.cell(5, 2))

        table.cell(6, 0).merge(table.cell(6, 1))
        table.cell(6, 0).merge(table.cell(6, 2))
        table.cell(7, 0).merge(table.cell(7, 1))
        table.cell(7, 0).merge(table.cell(7, 2))

        table.cell(8, 0).merge(table.cell(8, 1))
        table.cell(8, 0).merge(table.cell(8, 2))
        table.cell(9, 0).merge(table.cell(9, 1))
        table.cell(9, 0).merge(table.cell(9, 2))

        row = table.rows[0]
        row.cells[0].text = '漏洞ID'
        row.cells[1].text = '软件包名'
        row.cells[2].text = '评分'
        for cell in row.cells:
            color_for_cell(cell)
            paragraph = cell.paragraphs[0]
            # 设置字体样式
            run = paragraph.runs[0]
            run.font.bold = True
            # run.font.color.rgb = RGBColor(0,0,0)

        row = table.rows[1]
        row.cells[0].text = vul_info.cveid
        row.cells[1].text = vul_info.source
        if vul_info.score == '-1':
            row.cells[2].text = 'None' + ' (' + vul_info.level + ')'
        else:
            row.cells[2].text = vul_info.score + ' (' + vul_info.level + ')'
        run = row.cells[2].paragraphs[0].runs[0]
        if float(vul_info.score) >= 9:
            # 设置颜色（RGBColor(r, g, b) 中 r、g、b 分别表示红、绿、蓝的值）
            run.font.color.rgb = RGBColor(192, 0, 0)
        elif float(vul_info.score) >= 7:
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif float(vul_info.score) >= 4:
            run.font.color.rgb = RGBColor(247, 150, 70)
        else:
            run.font.color.rgb = RGBColor(255, 192, 0)

        table.cell(2, 0).text = "修复状态"
        color_for_cell(table.cell(2, 0))
        paragraph = table.cell(2, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(3, 0).text = "已修复" if vul_info.is_fixed else "未修复"
        paragraph = table.cell(3, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        if not vul_info.is_fixed:
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run.font.color.rgb = RGBColor(146, 208, 80)

        table.cell(4, 0).text = "漏洞描述"
        color_for_cell(table.cell(4, 0))
        paragraph = table.cell(4, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(5, 0).text = vul_info.description

        table.cell(6, 0).text = "修复建议"
        color_for_cell(table.cell(6, 0))
        paragraph = table.cell(6, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(7, 0).text = vul_info.fix_suggest

        table.cell(8, 0).text = "参考链接"
        color_for_cell(table.cell(8, 0))
        paragraph = table.cell(8, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(9, 0).text = vul_info.reference_link

        table.rows[1].height = Inches(0.3)
        table.rows[3].height = Inches(0.3)
        table.rows[5].height = Inches(0.3)
        table.rows[7].height = Inches(0.3)
        table.rows[9].height = Inches(0.3)

        doc_element = index._element
        doc_element.addprevious(heading._element)
        doc_element.addprevious(table._element)
    else:
        table = doc.add_table(12, 3, style="Table Grid")
        table.autofit = False
        table.allow_autofit = False
        # 设置每列的宽度，单位为英寸
        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(2)
        table.columns[2].width = Inches(2)

        table.cell(2, 0).merge(table.cell(2, 1))
        table.cell(2, 0).merge(table.cell(2, 2))
        table.cell(3, 0).merge(table.cell(3, 1))
        table.cell(3, 0).merge(table.cell(3, 2))

        table.cell(4, 0).merge(table.cell(4, 1))
        table.cell(4, 0).merge(table.cell(4, 2))
        table.cell(5, 0).merge(table.cell(5, 1))
        table.cell(5, 0).merge(table.cell(5, 2))

        table.cell(6, 0).merge(table.cell(6, 1))
        table.cell(6, 0).merge(table.cell(6, 2))
        table.cell(7, 0).merge(table.cell(7, 1))
        table.cell(7, 0).merge(table.cell(7, 2))

        table.cell(8, 0).merge(table.cell(8, 1))
        table.cell(8, 0).merge(table.cell(8, 2))
        table.cell(9, 0).merge(table.cell(9, 1))
        table.cell(9, 0).merge(table.cell(9, 2))

        table.cell(10, 0).merge(table.cell(10, 1))
        table.cell(10, 0).merge(table.cell(10, 2))
        table.cell(11, 0).merge(table.cell(11, 1))
        table.cell(11, 0).merge(table.cell(11, 2))

        row = table.rows[0]
        row.cells[0].text = 'CVE ID'
        row.cells[1].text = '软件包名'
        row.cells[2].text = '评分'
        for cell in row.cells:
            color_for_cell(cell)
            paragraph = cell.paragraphs[0]
            # 设置字体样式
            run = paragraph.runs[0]
            run.font.bold = True
            # run.font.color.rgb = RGBColor(255, 255, 255)

        row = table.rows[1]
        row.cells[0].text = vul_info.cveid
        row.cells[1].text = vul_info.source
        if vul_info.score == '-1':
            row.cells[2].text = 'None' + ' (' + vul_info.level + ')'
        else:
            row.cells[2].text = vul_info.score + ' (' + vul_info.level + ')'
        run = row.cells[2].paragraphs[0].runs[0]
        if float(vul_info.score) >= 9:
            # 设置颜色（RGBColor(r, g, b) 中 r、g、b 分别表示红、绿、蓝的值）
            run.font.color.rgb = RGBColor(192, 0, 0)
        elif float(vul_info.score) >= 7:
            run.font.color.rgb = RGBColor(255, 0, 0)
        elif float(vul_info.score) >= 4:
            run.font.color.rgb = RGBColor(247, 150, 70)
        else:
            run.font.color.rgb = RGBColor(255, 192, 0)

        table.cell(2, 0).text = "修复状态"
        color_for_cell(table.cell(2, 0))
        paragraph = table.cell(2, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)
        table.cell(3, 0).text = "已修复" if vul_info.is_fixed else "未修复"
        paragraph = table.cell(3, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        if not vul_info.is_fixed:
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run.font.color.rgb = RGBColor(146, 208, 80)

        table.cell(4, 0).text = "主机受影响软件包"
        color_for_cell(table.cell(4, 0))
        paragraph = table.cell(4, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)
        table.cell(5, 0).text = vul_info.affectPkg

        table.cell(6, 0).text = "漏洞描述"
        color_for_cell(table.cell(6, 0))
        paragraph = table.cell(6, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(7, 0).text = vul_info.description

        table.cell(8, 0).text = "修复建议"
        color_for_cell(table.cell(8, 0))
        paragraph = table.cell(8, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(9, 0).text = vul_info.fix_suggest

        table.cell(10, 0).text = "参考链接"
        color_for_cell(table.cell(10, 0))
        paragraph = table.cell(10, 0).paragraphs[0]
        # 设置字体样式
        run = paragraph.runs[0]
        run.font.bold = True
        # run.font.color.rgb = RGBColor(255, 255, 255)

        table.cell(11, 0).text = vul_info.reference_link

        table.rows[1].height = Inches(0.3)
        table.rows[3].height = Inches(0.3)
        table.rows[5].height = Inches(0.3)
        table.rows[7].height = Inches(0.3)
        table.rows[9].height = Inches(0.3)
        table.rows[11].height = Inches(0.3)

        doc_element = index._element
        doc_element.addprevious(heading._element)
        doc_element.addprevious(table._element)


def get_origin_vul_data(exist_vul):
    list1 = []
    list2 = []
    list3 = []
    list4 = []
    for vul in exist_vul:
        if vul['score'] is None or vul['score'] == '':
            vul['score'] = '-1'
        if 'is_fixed' not in vul:
            vul['is_fixed'] = False
        if 'affectPkg' in vul:
            temp = VulInfo(vul['cveid'], vul['source'], vul['score'], vul['level'], vul['cve_description'],
                           vul['is_fixed'], vul['fix_suggest'], vul['reference_link'],
                           vul['affectPkg'])
        else:
            temp = VulInfo(vul['cveid'], vul['source'], vul['score'], vul['level'], vul['cve_description'],
                           vul['is_fixed'], vul['fix_suggest'], vul['reference_link'])
        if float(temp.score) < 4.0:
            list4.append(temp)
        elif float(temp.score) < 7.0:
            list3.append(temp)
        elif float(temp.score) < 9.0:
            list2.append(temp)
        else:
            list1.append(temp)

    return list1, list2, list3, list4


def set_single_spacing(paragraph):
    # 设置段落的行距为单倍行距
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def analyze_year_info(exist_vul):
    year_dict = dict()
    for vul in exist_vul:
        year_info = vul['cveid'].split('-')[1]
        if year_info in year_dict:
            year_dict[year_info] += 1
        else:
            year_dict[year_info] = 1
    year_sorted = {k: v for k, v in sorted(year_dict.items(), key=lambda item: item[1], reverse=False)}
    plt.bar(year_sorted.keys(), height=year_sorted.values(), width=0.3)
    plt.xlabel('Year')
    plt.ylabel('Number')
    image_stream = BytesIO()
    plt.savefig(image_stream, format='png')
    plt.close()
    return image_stream


def generate_doc(exist_vul, username, hostname, time_ob, local_version, pkg_num, server_type):
    list1, list2, list3, list4 = get_origin_vul_data(exist_vul)
    fixed_num = 0
    for vul in exist_vul:
        if 'is_fixed' in vul:
            if vul['is_fixed']:
                fixed_num += 1

    analyze_year_info(exist_vul)
    # 生成饼图
    labels = ['严重', '高危', '中危', '低危']
    sizes = [len(list1), len(list2), len(list3), len(list4)]
    colors = [(192 / 255, 0, 0), (255 / 255, 0, 0), (247 / 255, 150 / 255, 70 / 255), (255 / 255, 192 / 255, 0)]

    fig, ax = plt.subplots()
    ax.pie(sizes, colors=colors, autopct='%1.1f%%', startangle=90)
    legend_patches = [Patch(color=color, label=label) for color, label in zip(colors, labels)]
    plt.legend(handles=legend_patches, bbox_to_anchor=(1.1, 1.05), loc="upper right")
    ax.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.

    # 将图像保存到字节流
    image_stream = BytesIO()
    plt.savefig(image_stream, format='png')
    plt.close()

    # 生成柱状图
    plt.bar(labels, sizes, color=colors)
    plt.xlabel('漏洞风险等级')
    plt.ylabel('数量')
    image_stream_bar = BytesIO()
    plt.savefig(image_stream_bar, format='png')
    plt.close()

    base_index = 15

    # Initialize a Document
    doc = Document(global_variable.DOC_TMP_PATH + '/doc_generator/new_file.docx')
    index = doc.paragraphs[base_index + 1]
    index.insert_paragraph_before(
        "本次扫描共扫描软件包" + str(pkg_num) + "个，目标主机【" + hostname + "】共检测出漏洞" + str(
            len(exist_vul)) + "个，其中严重漏洞" + str(len(
            list1)) + "个，高危漏洞" + str(len(list2)) + "个，中危漏洞" + str(len(list3)) + "个，低危漏洞" + str(len(
            list4)) + "个。修复漏洞" + str(fixed_num) + "个（漏洞风险等级的分类请参阅目录2 参考标准）。")

    table0 = doc.tables[0]
    table0.cell(3, 3).paragraphs[0].text = hostname
    table0.cell(4, 0).paragraphs[0].text = username
    table0.cell(4, 2).paragraphs[0].text = time_ob[2] + '秒'
    table0.cell(5, 1).paragraphs[0].text = time_ob[0]
    table0.cell(6, 0).paragraphs[0].text = time_ob[1]
    table0.cell(6, 3).paragraphs[0].text = "正常"
    table0.cell(7, 2).paragraphs[0].text = local_version

    table1 = doc.tables[1]
    table1.cell(0, 1).paragraphs[0].runs[0].text = str(len(list1))
    table1.cell(0, 2).paragraphs[0].runs[0].text = str(len(list2))
    table1.cell(1, 0).paragraphs[0].runs[0].text = str(len(list3))
    table1.cell(1, 2).paragraphs[0].runs[0].text = str(len(list4))

    # todo
    table1.cell(1, 3).paragraphs[0].runs[0].text = str(fixed_num)

    pic_index = doc.paragraphs[base_index + 16]
    temp = pic_index.insert_paragraph_before("").add_run()
    temp.add_picture(image_stream, width=Inches(4.5))
    set_single_spacing(doc.paragraphs[base_index + 16])

    pic_index = doc.paragraphs[base_index + 18]
    temp = pic_index.insert_paragraph_before("").add_run()
    temp.add_picture(image_stream_bar, width=Inches(4.5))
    set_single_spacing(doc.paragraphs[base_index + 18])

    index1 = doc.paragraphs[base_index + 22]
    index2 = doc.paragraphs[base_index + 23]
    index3 = doc.paragraphs[base_index + 24]
    index4 = doc.paragraphs[base_index + 25]

    for vul in list1:
        insert_table(doc, vul, index1, server_type)
    for vul in list2:
        insert_table(doc, vul, index2, server_type)
    for vul in list3:
        insert_table(doc, vul, index3, server_type)
    for vul in list4:
        insert_table(doc, vul, index4, server_type)
    filename = hostname + ' ' + time_ob[0] + '.docx'
    doc.save(global_variable.DOC_FILE_PATH + '/word/' + filename)
    return filename
